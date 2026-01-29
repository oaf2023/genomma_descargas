from docx import Document
from docx.shared import Pt
import os

def create_documentation():
    doc = Document()

    # Title
    doc.add_heading('Documentación Técnica de la Aplicación Genomma Lab', 0)

    # Overview
    doc.add_heading('1. Visión General', level=1)
    doc.add_paragraph(
        'Esta aplicación es una solución integral para la extracción, transformación y carga (ETL) '
        'de datos desde SQL Server hacia Snowflake, con una interfaz de visualización y consulta construida en Streamlit. '
        'Soporta operaciones multi-país (Chile, Colombia, Ecuador, Perú) y maneja reportes complejos y procedimientos almacenados.'
    )

    # Main Application
    doc.add_heading('2. Aplicación Principal (Frontend)', level=1)
    
    doc.add_heading('streamlit_app.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Punto de entrada principal para el dashboard de usuario. Permite la autenticación, navegación y visualización de datos almacenados en Snowflake.')
    
    p = doc.add_paragraph()
    p.add_run('Conexiones:').bold = True
    doc.add_paragraph('• Importa y utiliza "app_reportes_sql.py" para lógica de reportes.', style='List Bullet')
    doc.add_paragraph('• Conecta directamente con Snowflake usando credenciales de st.secrets o variables de entorno.', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Forma de Uso: ').bold = True
    p.add_run('Se ejecuta mediante el comando "streamlit run streamlit_app.py". Es la interfaz que ve el usuario final.')

    # Logic & Reports
    doc.add_heading('3. Lógica de Reportes y SQL', level=1)
    
    doc.add_heading('app_reportes_sql.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Contiene la lógica de negocio para la ejecución de procedimientos almacenados en SQL Server y la generación de reportes específicos por país.')
    
    p = doc.add_paragraph()
    p.add_run('Conexiones:').bold = True
    doc.add_paragraph('• Utiliza "pyodbc" para conectar con servidores SQL Server legacy.', style='List Bullet')
    doc.add_paragraph('• Interactúa con módulos de metadatos ("tabla_metadata") y hashing ("hash_control") si están disponibles.', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Forma de Uso: ').bold = True
    p.add_run('No se ejecuta directamente. Es importado como módulo por la aplicación principal para proveer funciones de backend.')

    # ETL Orchestration
    doc.add_heading('4. Orquestación ETL', level=1)
    
    doc.add_heading('pipeline_maestro.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Script maestro que coordina la ejecución secuencial de los pasos del proceso ETL.')
    
    p = doc.add_paragraph()
    p.add_run('Flujo de Trabajo (Pasos):').bold = True
    doc.add_paragraph('1. Verificación de configuración.', style='List Number')
    doc.add_paragraph('2. Normalización de headers (llama a script 2).', style='List Number')
    doc.add_paragraph('3. Renombrado de archivos (llama a script 3).', style='List Number')
    doc.add_paragraph('4. Carga a Snowflake (llama a script 4).', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('Forma de Uso: ').bold = True
    p.add_run('Se ejecuta manualmente o programado vía CLI: "python pipeline_maestro.py". Admite argumentos como "--dry-run" o "--step".')

    # ETL Scripts Details
    doc.add_heading('5. Scripts de Componentes ETL (Carpeta /etl)', level=1)
    
    # Script 1
    doc.add_heading('etl/1_descargar_sql_server.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Aplicación Streamlit independiente diseñada para descargar datos masivos de SQL Server y guardarlos localmente o en Google Drive.')
    p = doc.add_paragraph()
    p.add_run('Conexión: ').bold = True
    p.add_run('Directa a SQL Server. Genera archivos CSV.')
    p = doc.add_paragraph()
    p.add_run('Uso: ').bold = True
    p.add_run('Puede ejecutarse interactivamente: "streamlit run etl/1_descargar_sql_server.py".')

    # Script 2
    doc.add_heading('etl/2_normalizar_headers.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Estandariza los nombres de las columnas en los archivos CSV descargados para asegurar consistencia antes de la carga.')
    p = doc.add_paragraph()
    p.add_run('Uso: ').bold = True
    p.add_run('Invocado por "pipeline_maestro.py" o individualmente.')

    # Script 3
    doc.add_heading('etl/3_renombrar_archivos.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Renombra los archivos CSV siguiendo convenciones de nombrado específicas requeridas para la ingesta en Snowflake.')
    p = doc.add_paragraph()
    p.add_run('Uso: ').bold = True
    p.add_run('Invocado por "pipeline_maestro.py".')

    # Script 4
    doc.add_heading('etl/4_cargar_snowflake.py', level=2)
    p = doc.add_paragraph()
    p.add_run('Función: ').bold = True
    p.add_run('Toma los archivos CSV procesados y realiza la carga (COPY) hacia las tablas staging o finales en Snowflake.')
    p = doc.add_paragraph()
    p.add_run('Conexión: ').bold = True
    p.add_run('Conecta con Snowflake usando el driver de Python.')
    p = doc.add_paragraph()
    p.add_run('Uso: ').bold = True
    p.add_run('Invocado al final del proceso por "pipeline_maestro.py".')

    # Save
    filename = 'Documentacion_Tecnica_Genomma_Lab.docx'
    doc.save(filename)
    print(f"Documento creado exitosamente: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_documentation()
